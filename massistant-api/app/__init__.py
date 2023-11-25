import mimetypes
import os
import uuid
from typing import Annotated, Union
import base64

import aiofiles
from celery.utils.log import logging
from fastapi import Depends, FastAPI, File, UploadFile, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session

from app.database import get_db
from app.queries import create_konspekt_upload, delete_konspekt_by_id, get_konspekt_by_id, get_konspekts
from app.response_models import (
    DefaultErrorResponse,
    KonspektsListItem,
    KonspektUploadSuccessResponse,
    TranscribeUploadSuccessResponse,
)
from app.request_models import UpdateKonspektText, UploadExtractedSummary, UploadExtractedTerms, UploadTranscribedText

from app.tasks import send_summary_task, send_terms_task, send_transcribe_task, celeryFeedback

import docx

import tempfile

app = FastAPI()

# КОСТЫЛИ ДЛЯ ЛОКАЛЬНОГО ОКРУЖЕНИЯ
origins = [
    "http://massistant.local",
    "https://massistant.local",
    "http://localhost",
    "https://localhost",
    "http://localhost:8080",
    "https://localhost:8080",
    "http://localhost:5173",
    "https://localhost:5173",
    "http://localhost:8000",
    "https://localhost:8000",
    "http://127.0.0.1",
    "https://127.0.0.1",
    "http://127.0.0.1:8080",
    "https://127.0.0.1:8080",
    "http://127.0.0.1:5173",
    "https://127.0.0.1:5173",
    "http://127.0.0.1:8000",
    "https://127.0.0.1:8000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

uploads_dir = os.path.abspath("./uploads/")

allowed_mimetypes = [
    "audio/mpeg"
]


@app.get("/")
def read_root():
    return {"Hello": "World"}


@app.get("/items/{item_id}")
def read_item(item_id: int, q: Union[str, None] = None):
    return {"item_id": item_id, "q": q}


@app.post("/api/konspekt")
async def upload_konspekt(mode: Union[None, str], audio: UploadFile, db: Session = Depends(get_db)) -> Union[
        KonspektUploadSuccessResponse, DefaultErrorResponse]:
    filename = audio.filename
    og_filename = audio.filename
    mime = audio.content_type

    if mime not in allowed_mimetypes:
        return DefaultErrorResponse(
            error_msg="Тип файла не поддерживается",
            error_key="upload.unsupported_file"
        )

    ext = mimetypes.guess_extension(mime)
    if ext is None:
        return DefaultErrorResponse(
            error_msg="Неизвестный тип файла",
            error_key="upload.unknown_filetype"
        )

    filename = str(uuid.uuid4()) + ext

    path_to_save = os.path.join(uploads_dir, filename)
    async with aiofiles.open(path_to_save, 'wb') as out_file:
        while content := await audio.read(1024):  # async read chunk
            await out_file.write(content)  # async write chunk

    new_upload = create_konspekt_upload(
        db,
        orignal_filename=og_filename,
        filename=filename
    )
    print(new_upload.id)

    if mode != "fast" and mode != "precise":
        mode = "fast"

    send_transcribe_task(new_upload.id, filename, mode)

    return KonspektUploadSuccessResponse(
        msg="Конспект загружен",
        key="upload.success",
        filename=filename
    )


# TODO: тут ещё будут роуты для транскрибации и тп, но пока хз как они выглядят

@app.get("/api/konspekt")
async def get_all_konspekts(db: Session = Depends(get_db)) -> list[KonspektsListItem]:
    konspekts: list[KonspektsListItem] = []

    for k in get_konspekts(db):
        konspekts.append(KonspektsListItem(
            id=k.id,
            created_at=k.created_at,
            original_filename=k.original_filename,
            filename=k.filename,
            status=k.status,
            trans_text=k.transcribe,
            summary=k.summary,
            glossary=k.glossary
        ))

    return konspekts


@app.get("/api/konspekt/audio/{filename}")
def get_konspekt_file(filename: str):
    filepath = os.path.join(uploads_dir, filename)
    if not os.path.exists(filepath):
        return DefaultErrorResponse(
            error_msg="Данный файл отсутствует",
            error_key="upload.missing_file"
        )

    ext = mimetypes.guess_type(filepath)

    def iterfile():
        with open(filepath, "rb") as f:
            yield from f

    headers = {'Content-Disposition': f'attachment; filename="{filename}"'}
    return StreamingResponse(iterfile(), headers=headers, media_type=ext[0])

# @celeryFeedback.task
# def set_transcribed_text(konspekt_id: int, text: str):
#     print("Celery feedback: Got transcribed text for id =", konspekt_id)
#
#     db = get_db()
#     if db is None:
#         logging.er
#
#     konspekt_upload = get_konspekt_by_id(db, konspekt_id)
#
#     if konspekt_upload is None:
#         return DefaultErrorResponse(
#             error_msg="Неверный ID конспекта",
#             error_key="transcribe_upload.invalid_konspekt_id"
#         )


@app.post("/api/konspekt/{kid}/transcribed")
def set_konspket_transcribed_text(kid: int, transcribed: UploadTranscribedText, db: Session = Depends(get_db)):
    print("Got transcribed text for id =", kid)

    konspekt_upload = get_konspekt_by_id(db, kid)

    if konspekt_upload is None:
        return DefaultErrorResponse(
            error_msg="Неверный ID конспекта",
            error_key="transcribe_upload.invalid_konspekt_id"
        )

    text = base64.b64decode(transcribed.data).decode("utf8")
    konspekt_upload.transcribe = text
    konspekt_upload.status = "transcribed"
    db.commit()

    send_terms_task(kid, text)
    send_summary_task(kid, text)

    return TranscribeUploadSuccessResponse(
        msg="Транскрибированный текст сохранен в базу данных",
        key="transcribe_upload.success"
    )


@app.post("/api/konspekt/{kid}/terms")
def set_konspket_extracted_terms(kid: int, terms: UploadExtractedTerms, db: Session = Depends(get_db)):
    print("Got terms text for id =", kid)

    konspekt_upload = get_konspekt_by_id(db, kid)

    if konspekt_upload is None:
        return DefaultErrorResponse(
            error_msg="Неверный ID конспекта",
            error_key="terms.invalid_konspekt_id"
        )

    konspekt_upload.glossary = {
        'terms': terms.data  # 🤡
    }
    db.commit()

    return TranscribeUploadSuccessResponse(
        msg="Термины сохранены в базу данных",
        key="terms.success"
    )


@app.post("/api/konspekt/{kid}/summary")
def set_konspket_extracted_summary(kid: int, summary: UploadExtractedSummary, db: Session = Depends(get_db)):
    print("Got terms text for id =", kid)

    konspekt_upload = get_konspekt_by_id(db, kid)

    if konspekt_upload is None:
        return DefaultErrorResponse(
            error_msg="Неверный ID конспекта",
            error_key="summary.invalid_konspekt_id"
        )

    konspekt_upload.summary = summary.data
    db.commit()

    return TranscribeUploadSuccessResponse(
        msg="Суммаризация сохранена в базу данных",
        key="summary.success"
    )


@app.get("/api/konspekt/{kid}/download")
def download_konspekt(kid: int, db: Session = Depends(get_db)):
    konspekt_upload = get_konspekt_by_id(db, kid)

    if konspekt_upload is None:
        return DefaultErrorResponse(
            error_msg="Неверный ID конспекта",
            error_key="download.invalid_konspekt_id"
        )

    document = docx.Document()
    document.add_heading(konspekt_upload.original_filename, 0)
    document.add_paragraph("## Транскрибация\n", style='Heading 1')
    document.add_paragraph("None" if konspekt_upload.transcribe is None else konspekt_upload.transcribe)
    document.add_paragraph("## Глоссарий\n", style='Heading 1')

    terms = konspekt_upload.glossary['terms']
    for i in range(len(terms)):
        document.add_paragraph(f"{terms[i]['term']} - {terms[i]['meaning']}")
    document.add_paragraph("## Краткое содержание\n", style='Heading 1')
    if konspekt_upload.summary is None:
        document.add_paragraph("None")
    else:
        document.add_paragraph("###Введение\n", style='Heading 2')
        document.add_paragraph("None" if konspekt_upload.summary['introduction'] is None else konspekt_upload.summary['introduction'])
        document.add_paragraph("###Основная часть\n", style='Heading 2')
        document.add_paragraph("None" if konspekt_upload.summary['main_part'] is None else konspekt_upload.summary['main_part'])
        document.add_paragraph("###Основная часть\n", style='Heading 2')
        document.add_paragraph("None" if konspekt_upload.summary['conclusion'] is None else konspekt_upload.summary['conclusion'])

    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        document.save(tmp.name)
        return FileResponse(tmp.name, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.delete("/api/konspekt/{kid}")
def delete_konspekt(kid: int, db: Session = Depends(get_db)):
    delete_konspekt_by_id(db, kid)

    return TranscribeUploadSuccessResponse(
        msg="Конспект удален",
        key="delete.success"
    )


@app.post("/api/konspekt/{kid}/text")
def update_konspekt_text(kid: int, data: UpdateKonspektText, db: Session = Depends(get_db)):
    konspekt_upload = get_konspekt_by_id(db, kid)

    if konspekt_upload is None:
        return DefaultErrorResponse(
            error_msg="Неверный ID конспекта",
            error_key="update_text.invalid_konspekt_id"
        )

    konspekt_upload.transcribe = data.text
    db.commit()

    return TranscribeUploadSuccessResponse(
        msg="Конспект обновлен",
        key="update_text.success"
    )
